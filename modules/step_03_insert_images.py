"""
DOCX Image Swap — step_03_insert_images
=========================================
Reinject images into DOCX at anchor positions {{img_N}}.

1. Load working.docx from temp/
2. Load image_map.json from temp/
3. Find paragraphs with {{img_N}} anchors
4. Replace anchors with images from images/
5. Save result

Features:
- Preserves aspect ratio
- Limits max width/height
- Centers images
- If file not found — keeps anchor as text
"""

from pathlib import Path
import json
import re
import time
from typing import Dict, Any, Callable, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


def run(work_path: Path, step_config: Dict[str, Any], log_func: Callable[[str], None], ai_client=None, **kwargs) -> bool:
    """
    Insert images from images/ at anchor positions {{img_N}} in working.docx.

    Args:
        work_path: project working directory
        step_config: step parameters from config
        log_func: logging callback

    Returns:
        bool: True on success
    """
    start_time = time.time()

    if not HAS_DOCX:
        log_func("Missing library: python-docx")
        log_func("   Install: pip install python-docx")
        return False

    try:
        # --- PARAMETERS ---
        config = step_config.get("config", step_config)
        max_width_inches = config.get("max_width_inches", 5.5)
        max_height_inches = config.get("max_height_inches", 8.0)
        center_images = config.get("center_images", True)
        preserve_aspect = config.get("preserve_aspect_ratio", True)

        log_func("")
        log_func("=" * 70)
        log_func("   REINSERTING IMAGES INTO DOCUMENT")
        log_func("=" * 70)

        # --- PATHS ---
        temp_dir = work_path / "temp"
        images_dir = work_path / "images"
        working_docx = temp_dir / "working.docx"
        image_map_path = temp_dir / "image_map.json"

        # --- CHECKS ---
        if not working_docx.exists():
            log_func("Error: working.docx not found!")
            return False

        if not image_map_path.exists():
            log_func("Error: image_map.json not found!")
            log_func("   Run the extraction step first.")
            return False

        if not images_dir.exists():
            log_func("Error: images/ folder not found!")
            return False

        # --- LOAD MAP ---
        log_func("")
        log_func("Loading image map...")

        with open(image_map_path, "r", encoding="utf-8") as f:
            raw_map = json.load(f)

        # Normalize map: anchor -> filepath
        image_map: Dict[str, Path] = {}
        for anchor, filename in raw_map.items():
            if isinstance(filename, dict):
                filename = filename.get("filename", "")
            if not filename:
                continue

            file_path = images_dir / filename
            if file_path.exists():
                image_map[anchor] = file_path
            else:
                log_func(f"   Warning: file not found: {filename}")

        log_func(f"   Map entries: {len(raw_map)}")
        log_func(f"   Files found: {len(image_map)}")

        if not image_map:
            log_func("Warning: no images available for insertion")
            log_func("   Check the images/ folder")
            return True

        # --- LOAD DOCUMENT ---
        log_func("")
        log_func("Loading document...")
        doc = Document(working_docx)
        total_paragraphs = len(doc.paragraphs)
        log_func(f"   Paragraphs: {total_paragraphs}")

        # --- FIND AND REPLACE ANCHORS ---
        log_func("")
        log_func("Finding and replacing anchors...")
        log_func("-" * 70)

        anchor_pattern = re.compile(r'\{\{img_\d+\}\}')

        stats = {
            "found": 0,
            "replaced": 0,
            "not_in_map": 0,
            "errors": 0,
            "kept_as_text": 0,
        }

        # Collect paragraphs with anchors
        paragraphs_with_anchors = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if anchor_pattern.search(text):
                paragraphs_with_anchors.append((i, para, text))

        log_func(f"   Paragraphs with anchors: {len(paragraphs_with_anchors)}")

        for para_idx, para, text in paragraphs_with_anchors:
            anchors_in_para = anchor_pattern.findall(text)

            for anchor in anchors_in_para:
                stats["found"] += 1

                if anchor not in image_map:
                    stats["not_in_map"] += 1
                    log_func(f"   Warning: [{para_idx}] {anchor} — not in map, kept as text")
                    stats["kept_as_text"] += 1
                    continue

                image_path = image_map[anchor]

                try:
                    width_inches, height_inches = _calculate_image_size(
                        image_path, max_width_inches, max_height_inches,
                        preserve_aspect, log_func
                    )

                    only_anchor = text.strip() == anchor

                    if only_anchor:
                        _replace_paragraph_with_image(
                            para, image_path, width_inches, center_images
                        )
                    else:
                        _insert_image_after_paragraph(
                            doc, para, image_path, width_inches, center_images, anchor
                        )

                    size_kb = image_path.stat().st_size / 1024
                    log_func(f"   [{para_idx}] {anchor} -> {image_path.name} ({size_kb:.1f} KB, {width_inches:.1f}\")")
                    stats["replaced"] += 1

                except Exception as e:
                    log_func(f"   Error: [{para_idx}] {anchor} — {e}")
                    stats["errors"] += 1

        # --- SAVE ---
        log_func("")
        log_func("Saving document...")
        doc.save(working_docx)

        file_size_kb = working_docx.stat().st_size / 1024
        duration = time.time() - start_time

        # --- SUMMARY ---
        log_func("")
        log_func("=" * 70)
        log_func("INSERTION RESULTS")
        log_func("=" * 70)
        log_func(f"   Anchors found: {stats['found']}")
        log_func(f"   Images inserted: {stats['replaced']}")

        if stats["not_in_map"] > 0:
            log_func(f"   Not in map: {stats['not_in_map']}")
        if stats["kept_as_text"] > 0:
            log_func(f"   Kept as text: {stats['kept_as_text']}")
        if stats["errors"] > 0:
            log_func(f"   Errors: {stats['errors']}")

        log_func(f"   Document size: {file_size_kb:.1f} KB")
        log_func(f"   Duration: {duration:.1f} sec")
        log_func("")

        if stats["replaced"] > 0:
            log_func(f"Successfully inserted {stats['replaced']} images")
        else:
            log_func("No images were inserted")

        log_func("=" * 70)
        return True

    except Exception as e:
        log_func(f"Critical error: {e}")
        import traceback
        log_func("Traceback:")
        for line in traceback.format_exc().splitlines():
            if line.strip():
                log_func(f"   {line}")
        return False


# ═══════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════

def _calculate_image_size(
    image_path: Path,
    max_width: float,
    max_height: float,
    preserve_aspect: bool,
    log_func: Callable
) -> Tuple[float, float]:
    """Calculate image size for DOCX insertion. Returns (width_inches, height_inches)."""
    if not HAS_PIL:
        return (max_width, max_height)

    try:
        with Image.open(image_path) as img:
            w_px, h_px = img.size
            dpi = img.info.get('dpi', (96, 96))
            if isinstance(dpi, tuple):
                dpi_x = dpi[0] if dpi[0] > 0 else 96
                dpi_y = dpi[1] if dpi[1] > 0 else 96
            else:
                dpi_x = dpi_y = dpi if dpi > 0 else 96

            natural_w = w_px / dpi_x
            natural_h = h_px / dpi_y

            if preserve_aspect:
                scale_w = max_width / natural_w if natural_w > max_width else 1.0
                scale_h = max_height / natural_h if natural_h > max_height else 1.0
                scale = min(scale_w, scale_h)
                return (natural_w * scale, natural_h * scale)
            else:
                return (min(natural_w, max_width), min(natural_h, max_height))

    except Exception:
        return (max_width, max_height)


def _replace_paragraph_with_image(para, image_path: Path, width_inches: float, center: bool):
    """Replace paragraph content with an image. Paragraph stays, text replaced."""
    for run in para.runs:
        run.text = ""

    p_element = para._element
    children_to_remove = []
    for child in p_element:
        if child.tag != qn('w:pPr'):
            children_to_remove.append(child)
    for child in children_to_remove:
        p_element.remove(child)

    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _insert_image_after_paragraph(doc, para, image_path: Path, width_inches: float, center: bool, anchor: str):
    """Insert image in a new paragraph after the current one. Remove anchor from text."""
    for run in para.runs:
        if anchor in run.text:
            run.text = run.text.replace(anchor, "").strip()

    new_para = doc.add_paragraph()
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    if center:
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para._element.addnext(new_para._element)
